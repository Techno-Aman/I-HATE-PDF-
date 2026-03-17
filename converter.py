
import streamlit as st
from io import BytesIO
import zipfile
import tempfile
import os
from typing import List

# conversion libraries
from pdf2image import convert_from_bytes
import pdfplumber
from docx import Document
from pptx import Presentation
from pptx.util import Inches
from PIL import Image
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader

# ------------------------------
# Helper conversion functions
# ------------------------------
import streamlit as st

st.set_page_config(page_title="I💔PDF", layout="wide")

st.markdown("""
<style>

/* App background */
.stApp{
    background-color:black;
    color:white;
}

/* Card buttons */
div.stButton > button {
    background-color:#1f1f1f;
    color:white;
    border-radius:12px;
    padding:14px 18px;
    font-size:15px;
    border:2px solid transparent;
    transition: all 0.25s ease;
}

/* Hover glow */
div.stButton > button:hover {
    border:2px solid #00ADB5;
    box-shadow:0px 0px 12px #00ADB5;
}

/* Selected card */
.selected {
    background-color:white !important;
    color:black !important;
    border:2px solid #00ADB5 !important;
}

/* Convert button */
.convert button {
    background-color:#2a2a2a !important;
    color:white !important;
    border-radius:10px;
    padding:10px 18px;
}

.convert button:hover{
    border:2px solid #00ADB5;
    box-shadow:0px 0px 10px #00ADB5;
}

</style>
""", unsafe_allow_html=True)

st.markdown(
"""
<h1 style='text-align:center;'>I💔PDF</h1>
<p style='text-align:center;'>Convert PDF, Word, and PowerPoint easily</p>
""",
unsafe_allow_html=True
)

st.markdown("<br><br>", unsafe_allow_html=True)
if "conversion_type" not in st.session_state:
    st.session_state.conversion_type = None

col1, col2 = st.columns(2)

with col1:
    if st.button("📄 PDF → Word", use_container_width=True):
        st.session_state.conversion_type = "PDF TO WORD"

with col2:
     if st.button("📝 Word → PDF", use_container_width=True):
        st.session_state.conversion_type = "WORD TO PDF"
        
if st.session_state.conversion_type:
    col1, col2, col3 = st.columns([1,2,1])

    with col2:
        st.info(f"Selected conversion: {st.session_state.conversion_type}")



def pdf_to_docx(pdf_bytes: bytes) -> BytesIO:
    """Extract text (and some images) from PDF into a docx file."""
    doc = Document()
    with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                for para in text.split('\n'):
                    doc.add_paragraph(para)
            # try images on page
            try:
                images = page.images
                for i, img in enumerate(images):
                    # pdfplumber gives bbox; crop the page image
                    cropped = page.crop((img['x0'], img['top'], img['x1'], img['bottom'])).to_image(resolution=150)
                    im = cropped.original
                    img_byte_arr = BytesIO()
                    im.save(img_byte_arr, format='PNG')
                    img_byte_arr.seek(0)
                    doc.add_picture(img_byte_arr, width=Inches(4))
            except Exception:
                pass
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def docx_to_pdf(docx_bytes: bytes) -> BytesIO:
    """Create a simple PDF from docx text. This is a best-effort layout (no exact pagination).
    For high-fidelity conversion install LibreOffice and use it instead (instructions in README).
    """
    doc = Document(BytesIO(docx_bytes))
    out = BytesIO()
    c = canvas.Canvas(out, pagesize=letter)
    width, height = letter
    margin = 50
    y = height - margin
    line_height = 12
    for para in doc.paragraphs:
        text = para.text
        if not text.strip():
            y -= line_height
            if y < margin:
                c.showPage()
                y = height - margin
            continue
        # simple wrap
        words = text.split()
        line = ''
        for w in words:
            trial = (line + ' ' + w).strip()
            if c.stringWidth(trial) > (width - 2 * margin):
                c.drawString(margin, y, line)
                y -= line_height
                line = w
                if y < margin:
                    c.showPage()
                    y = height - margin
            else:
                line = trial
        if line:
            c.drawString(margin, y, line)
            y -= line_height
            if y < margin:
                c.showPage()
                y = height - margin
    c.save()
    out.seek(0)
    return out


# ------------------------------
# Streamlit UI
# ------------------------------

st.markdown("<br><br>", unsafe_allow_html=True)

colA, colB, colC = st.columns([1,2,1])

with colB:
    uploaded_files = st.file_uploader(
        "Upload your files",
        accept_multiple_files=True,
        type=["pdf","docx","pptx"]
    )

conversion = st.session_state.conversion_type

st.markdown('<div class="convert">', unsafe_allow_html=True)

colx, coly, colz = st.columns([2,1,2])
with coly:
    process_button = st.button("Convert File", use_container_width=True)

st.markdown('</div>', unsafe_allow_html=True)

# warning if card not selected
if process_button and not conversion:
    st.warning("Please select a conversion type first.")
if process_button and uploaded_files and conversion:

    results = []

    for up in uploaded_files:
        name = up.name
        data = up.read()

        filename = None
        out = None

        if conversion == "PDF TO WORD":
            out = pdf_to_docx(data)
            filename = name.replace(".pdf", ".docx")

        elif conversion == "WORD TO PDF":
            out = docx_to_pdf(data)
            filename = name.replace(".docx", ".pdf")

        if filename and out:
            results.append((filename, out.getvalue()))
    # SINGLE FILE DOWNLOAD
    if len(results) == 1:

        col1, col2, col3 = st.columns([1,2,1])

        with col2:
            download_clicked = st.download_button(
                "Download File",
                data=results[0][1],
                file_name=results[0][0]
            )

        if download_clicked:
            st.session_state.conversion_type = None

    # MULTIPLE FILES ZIP
    else:

        zip_buffer = BytesIO()

        with zipfile.ZipFile(zip_buffer, "w") as zf:
            for filename, data in results:
                 zf.writestr(filename, data)

        col1, col2, col3 = st.columns([1,2,1])

        with col2:
            download_clicked = st.download_button(
                "Download ZIP",
                data=zip_buffer.getvalue(),
                file_name="converted_files.zip"
            )

        if download_clicked:
            st.session_state.conversion_type = None